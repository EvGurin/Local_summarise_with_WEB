import os
import uuid
import threading
import queue
from flask import Flask, render_template, request, jsonify, send_from_directory, session
from werkzeug.utils import secure_filename
from flask_session import Session
from transformers import AutoTokenizer, AutoModelForCausalLM
from moviepy.video.io.VideoFileClip import VideoFileClip
import whisper
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
import torch
import time

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['RESULT_FOLDER'] = 'results'
app.config['SESSION_FOLDER'] = 'flask_session'  # По умолчанию Flask-Session хранит файлы здесь
app.config['ALLOWED_EXTENSIONS'] = {'mp4', 'avi', 'mov', 'mkv'}

app.secret_key = 'super_very_secret_key_here'
app.config['SESSION_TYPE'] = 'filesystem'
Session(app)

# Создаём каталоги, если их нет
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['RESULT_FOLDER'], exist_ok=True)
os.makedirs(app.config['SESSION_FOLDER'], exist_ok=True)

# Глобальный словарь задач
tasks = {}
# Глобальная очередь задач
task_queue = queue.Queue()

# Работаем только на CPU
device = 'cpu'
print(f"Используемое устройство: {device}")

# Загрузка модели Llama
model_path = "./model_32_16"
tokenizer = AutoTokenizer.from_pretrained(model_path)
model = AutoModelForCausalLM.from_pretrained(model_path).to(device)
model.eval()

# Загрузка Whisper на CPU
whisper_model = whisper.load_model("turbo", device="cpu")
whisper_model.eval()


def cleanup_flask_session():
    """
    Удаляем сессионные файлы в папке flask_session, которым более 24 часов.
    """
    now = time.time()
    session_folder = app.config['SESSION_FOLDER']
    if not os.path.exists(session_folder):
        return

    for filename in os.listdir(session_folder):
        file_path = os.path.join(session_folder, filename)
        if os.path.isfile(file_path):
            file_mtime = os.path.getmtime(file_path)
            # Проверяем возраст файла
            if now - file_mtime > 86400:  # старше суток
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Ошибка при удалении файла сессии {file_path}: {e}")


def cleanup_results_folder():
    """
    Удаляем ВСЕ файлы из results,
    у которых дата последнего изменения более суток.
    """
    now = time.time()
    for filename in os.listdir(app.config['RESULT_FOLDER']):
        file_path = os.path.join(app.config['RESULT_FOLDER'], filename)
        if os.path.isfile(file_path):
            mtime = os.path.getmtime(file_path)
            if now - mtime > 86400:  # 24 часа
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Ошибка при удалении файла из results: {file_path}: {e}")


def cleanup_files():
    """
    Удаляем файлы из tasks, которые завершены более суток назад,
    а также вызываем очистку всей папки results и папки flask_session.
    """
    now = time.time()
    to_delete = []
    for t_id, info in tasks.items():
        if info.get('status') == 'Завершено!':
            finish_time = info.get('finish_time')
            if finish_time and (now - finish_time) > 86400:  # Старше суток
                # Удаляем финальные файлы
                for f_key in ['summary_docx', 'summary_txt', 'transcription_txt']:
                    f_path = info.get(f_key)
                    if f_path and os.path.exists(f_path):
                        try:
                            os.remove(f_path)
                        except Exception as e:
                            print(f"Ошибка при удалении итогового файла {f_path}: {e}")
                to_delete.append(t_id)

    # Удаляем задачи из словаря
    for t_id in to_delete:
        del tasks[t_id]

    # Чистим папку results
    cleanup_results_folder()
    # Чистим папку flask_session
    cleanup_flask_session()


def video_to_mp3(video_path, output_path, task_id):
    if not os.path.exists(video_path):
        raise FileNotFoundError(f"Видео файл не найден: {video_path}")
    try:
        tasks[task_id]['status'] = 'Конвертация видео в аудио...'
        video = VideoFileClip(video_path)
        video.audio.write_audiofile(output_path)
        tasks[task_id]['status'] = 'Конвертация видео в аудио завершена.'
    except Exception as e:
        tasks[task_id]['status'] = f'Ошибка при конвертации видео в аудио: {str(e)}'
        raise e
    finally:
        # Закрываем VideoFileClip, чтобы освободить ресурсы
        video.close()


def make_with_whisper(audio_file, task_id):
    try:
        tasks[task_id]['status'] = 'Транскрибация аудио...'
        with torch.no_grad():
            result = whisper_model.transcribe(audio_file, language="ru")
        tasks[task_id]['status'] = 'Транскрибация аудио завершена.'
    except Exception as e:
        tasks[task_id]['status'] = f'Ошибка при транскрибации аудио: {str(e)}'
        raise e

    try:
        original_filename = tasks[task_id]['filename']
        # Добавляем также уникальный task_id, чтобы избежать коллизий
        transcription_filename = f"{task_id}_transcription.txt"
        transcription_path = os.path.join(app.config['RESULT_FOLDER'], transcription_filename)

        with open(transcription_path, "w", encoding="utf-8") as f:
            f.write(result['text'])
        tasks[task_id]['transcription_txt'] = transcription_path
        tasks[task_id]['status'] = 'Сохранение транскрибированного текста.'
    except Exception as e:
        tasks[task_id]['status'] = f'Ошибка при сохранении транскрибированного текста: {str(e)}'
        raise e

    return result['text']


def markdown_to_word(md_text, output_file, task_id):
    try:
        tasks[task_id]['status'] = 'Преобразование Markdown в DOCX...'
        output_file_docx = os.path.join(app.config['RESULT_FOLDER'], f'{task_id}_{output_file}.docx')
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, "html.parser")
        doc = Document()

        def set_text_style(run, font_size=14, color=(0, 0, 0)):
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(*color)

        def process_html(element, parent_paragraph=None):
            if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                level = int(element.name[1])
                text = element.get_text(strip=True)
                if text:
                    heading = doc.add_heading(text, level=level)
                    for run in heading.runs:
                        set_text_style(run, font_size=14 + (5 - level) * 2)
            elif element.name == "p":
                text = element.get_text(strip=True)
                if text:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(text)
                    set_text_style(run)
            elif element.name == "ul":
                for li in element.find_all("li", recursive=False):
                    process_html(li)
            elif element.name == "li":
                text = element.get_text(strip=True)
                if text:
                    paragraph = doc.add_paragraph(style="List Bullet")
                    for child in element.contents:
                        if isinstance(child, str):
                            run = paragraph.add_run(child.strip())
                            set_text_style(run)
                        elif child.name:
                            process_html(child, parent_paragraph=paragraph)
            elif element.name == "strong":
                if parent_paragraph:
                    text = element.get_text(strip=True)
                    if text:
                        run = parent_paragraph.add_run(text)
                        run.bold = True
                        set_text_style(run)
            elif element.name == "em":
                if parent_paragraph:
                    text = element.get_text(strip=True)
                    if text:
                        run = parent_paragraph.add_run(text)
                        run.italic = True
                        set_text_style(run)
            elif element.string:
                text = element.string.strip()
                if text and parent_paragraph:
                    run = parent_paragraph.add_run(text)
                    set_text_style(run)

        if soup.body:
            elements = soup.body.contents
        else:
            elements = soup.contents

        for element in elements:
            if isinstance(element, str):
                text = element.strip()
                if text:
                    paragraph = doc.add_paragraph(text)
                    for run in paragraph.runs:
                        set_text_style(run)
            elif element.name:
                process_html(element)

        doc.save(output_file_docx)
        tasks[task_id]['summary_docx'] = output_file_docx
        tasks[task_id]['status'] = 'Преобразование Markdown в DOCX завершено.'
    except Exception as e:
        tasks[task_id]['status'] = f'Ошибка при преобразовании Markdown в DOCX: {str(e)}'
        raise e


def summary_to_txt(md_text, output_file, task_id):
    try:
        tasks[task_id]['status'] = 'Создание текстового конспекта...'
        output_file_txt = os.path.join(app.config['RESULT_FOLDER'], f'{task_id}_{output_file}.txt')
        plain_text = md_text
        # Упрощённая чистка от Markdown-символов
        plain_text = (plain_text
                      .replace('# ', '')
                      .replace('#', '')
                      .replace('*', '')
                      .replace('**', '')
                      .replace('_', ''))

        with open(output_file_txt, 'w', encoding="utf-8") as file:
            file.write(plain_text)
        tasks[task_id]['summary_txt'] = output_file_txt
        tasks[task_id]['status'] = 'Создание текстового конспекта завершено.'
    except Exception as e:
        tasks[task_id]['status'] = f'Ошибка при создании текстового конспекта: {str(e)}'
        raise e


def make_outputs(md_text, task_id):
    original_filename = tasks[task_id]['filename']
    base_name = os.path.splitext(original_filename)[0]
    # Теперь добавляем task_id в имени выходных файлов
    output_summary_base = f"{base_name}_конспект"
    markdown_to_word(md_text, output_summary_base, task_id)
    summary_to_txt(md_text, output_summary_base, task_id)


def llama(meeting_text):
    prompt = (
        "### User:\n"
        "Произведи подробный конспект следующего текста, разделив его на ключевые части и указав самые важные моменты в каждой из них. "
        "Приводи только краткую информацию, исключая все второстепенные детали, и следи, чтобы структура была логичной и понятной.\n\n"
        f"Текст:\n{meeting_text}\n\n"
        "### Assistant:"
    )
    with torch.no_grad():
        inputs = tokenizer(prompt, return_tensors="pt").to(device)
        outputs = model.generate(
            **inputs,
            max_new_tokens=512,
            temperature=0.7,
            top_p=0.9,
            pad_token_id=tokenizer.eos_token_id,
            eos_token_id=tokenizer.eos_token_id,
            do_sample=True
        )

        result = tokenizer.decode(outputs[0], skip_special_tokens=True)
        summary = result[len(prompt):].strip()
    return summary


def process_video_to_summary(task_id, video_path):
    try:
        tasks[task_id]['status'] = 'Начало обработки...'

        # Конвертация в MP3
        mp3_path = os.path.splitext(video_path)[0] + '.mp3'
        video_to_mp3(video_path, mp3_path, task_id)

        # Транскрибация
        meeting_text = make_with_whisper(mp3_path, task_id)

        # Генерация конспекта
        tasks[task_id]['status'] = 'Генерация конспекта...'
        summary = llama(meeting_text)

        # Создание выходных файлов
        make_outputs(summary, task_id)

        tasks[task_id]['status'] = 'Завершено!'
        tasks[task_id]['finish_time'] = time.time()

    except Exception as e:
        tasks[task_id]['status'] = f'Ошибка: {str(e)}'
    finally:
        # Удаление исходных файлов из папки uploads (видео и mp3)
        if os.path.exists(video_path):
            try:
                os.remove(video_path)
            except Exception as e:
                print(f"Ошибка при удалении исходного видео {video_path}: {e}")

        if os.path.exists(mp3_path):
            try:
                os.remove(mp3_path)
            except Exception as e:
                print(f"Ошибка при удалении mp3-файла {mp3_path}: {e}")


def worker():
    while True:
        task = task_queue.get()
        if task is None:
            break
        task_id, video_path = task
        process_video_to_summary(task_id, video_path)
        task_queue.task_done()


thread = threading.Thread(target=worker, daemon=True)
thread.start()


def get_queue_position(task_id):
    q_list = list(task_queue.queue)
    position = 1
    for item in q_list:
        if item[0] == task_id:
            return position
        position += 1
    return None


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


@app.route('/')
def index():
    try:
        cleanup_files()
        return render_template('index.html')
    except Exception as e:
        return f"Произошла ошибка при загрузке страницы: {str(e)}", 500


@app.route('/upload', methods=['POST'])
def upload():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Нет файла в запросе.'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Файл не выбран.'}), 400

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)

            if 'session_id' not in session:
                session['session_id'] = str(uuid.uuid4())
            session_id = session['session_id']

            unique_id = str(uuid.uuid4())
            saved_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")

            try:
                file.save(saved_path)
            except Exception as e:
                return jsonify({'error': f'Не удалось сохранить файл: {str(e)}'}), 500

            tasks[unique_id] = {
                'status': 'В очереди',
                'filename': filename,
                'session_id': session_id
            }

            task_queue.put((unique_id, saved_path))
            position = get_queue_position(unique_id)
            tasks[unique_id]['queue_position'] = position

            return jsonify({'task_id': unique_id, 'queue_position': position}), 200
        else:
            return jsonify({'error': 'Недопустимый тип файла.'}), 400
    except Exception as e:
        return jsonify({'error': f'Произошла ошибка на сервере: {str(e)}'}), 500


@app.route('/tasks', methods=['GET'])
def get_tasks():
    try:
        cleanup_files()
        if 'session_id' not in session:
            return jsonify({'tasks': []}), 200
        session_id = session['session_id']
        task_list = []
        for t_id, info in tasks.items():
            if info.get('session_id') == session_id:
                task_info = {
                    'task_id': t_id,
                    'filename': info['filename'],
                    'status': info['status'],
                    'queue_position': info.get('queue_position', None),
                    'summary_docx': os.path.basename(info.get('summary_docx')) if info.get('summary_docx') else None,
                    'summary_txt': os.path.basename(info.get('summary_txt')) if info.get('summary_txt') else None,
                    'transcription_txt': os.path.basename(info.get('transcription_txt')) if info.get('transcription_txt') else None
                }
                task_list.append(task_info)
        return jsonify({'tasks': task_list}), 200
    except Exception as e:
        return jsonify({'error': f'Ошибка при получении списка задач: {str(e)}'}), 500


@app.route('/status/<task_id>', methods=['GET'])
def status_route(task_id):
    try:
        task = tasks.get(task_id)
        if not task:
            return jsonify({'status': 'invalid'}), 404

        if task['status'] == 'В очереди':
            position = get_queue_position(task_id)
            task['queue_position'] = position

        response = {'status': task['status']}
        if task.get('queue_position') is not None:
            response['queue_position'] = task['queue_position']
        return jsonify(response), 200
    except Exception as e:
        return jsonify({'error': f'Ошибка при получении статуса задачи: {str(e)}'}), 500


@app.route('/download/<task_id>/<file_type>', methods=['GET'])
def download(task_id, file_type):
    try:
        task = tasks.get(task_id)
        if not task or task.get('status') != 'Завершено!':
            return jsonify({'error': 'Файл не готов или задача не найдена.'}), 404

        if file_type == 'docx':
            target_filename = task.get('summary_docx')
        elif file_type == 'txt':
            target_filename = task.get('summary_txt')
        elif file_type == 'transcription':
            target_filename = task.get('transcription_txt')
        else:
            return jsonify({'error': 'Недопустимый тип файла.'}), 400

        if target_filename and os.path.exists(target_filename):
            return send_from_directory(app.config['RESULT_FOLDER'], os.path.basename(target_filename), as_attachment=True)
        else:
            return jsonify({'error': 'Файл не найден.'}), 404
    except Exception as e:
        return jsonify({'error': f'Ошибка при скачивании файла: {str(e)}'}), 500


if __name__ == '__main__':
    # Можно явно указать debug=False, чтобы исключить автоматические перезапуски
    app.run(host='0.0.0.0', port=5000, debug=False)
